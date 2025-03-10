import azure.functions as func 
import json
import base64
import io
import logging
from docx import Document
from docx.shared import Pt, RGBColor 
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls 
from docx.oxml import OxmlElement

app = func.FunctionApp(http_auth_level=func.AuthLevel.FUNCTION)

@app.route(route="receiveDoc")


def receiveDoc(req: func.HttpRequest) -> func.HttpResponse:
    """
    Handles incoming HTTP request:
    - Receives a DOCX file (base64) + an array of Q&A data.
    - Inserts a formatted table at the end of the document.
    - Returns the modified DOCX as base64.
    """
    try:
        req_body = req.get_json()
        input_base64 = req_body.get("base64", "")
        qa_data = req_body.get("qa_data", [])

        if not input_base64:
            return func.HttpResponse(json.dumps({"message": "No input base64 provided", "base64": None}),
                                     status_code=400, mimetype="application/json")

        if not qa_data:
            return func.HttpResponse(json.dumps({"message": "No Q&A data provided", "base64": None}),
                                     status_code=400, mimetype="application/json")

        # Decode base64 to DOCX bytes
        doc_bytes = base64.b64decode(input_base64)

        # Process document (Add table)
        modified_docx_bytes, message = process_docx(doc_bytes, qa_data)

        # If processing failed
        if modified_docx_bytes is None:
            return func.HttpResponse(json.dumps({"message": message, "base64": None}),
                                     status_code=500, mimetype="application/json")

        # Encode modified DOCX as base64
        output_base64 = base64.b64encode(modified_docx_bytes).decode("utf-8")

        return func.HttpResponse(json.dumps({"message": message, "base64": output_base64}),
                                 status_code=200, mimetype="application/json")

    except Exception as e:
        return func.HttpResponse(json.dumps({"message": f"Unexpected error: {str(e)}", "base64": None}),
                                 status_code=500, mimetype="application/json")

def process_docx(doc_bytes, qa_data):
    """
    Processes the DOCX document:
    - Inserts a new table just before the "Signature" section using the given Q&A data.
    - Supports the new format: [{"question": "...", "answer": "..."}].
    - Clears column 2 if "Section" is found before merging.
    - Returns modified DOCX as bytes.
    """
    try:
        # Load the DOCX from bytes
        doc = Document(io.BytesIO(doc_bytes))

        # Step 1: Find the "Signature" Paragraph
        signature_paragraph = None
        for para in doc.paragraphs:
            if "signature" in para.text.lower():  # Case-insensitive check
                signature_paragraph = para
                break  # Stop at the first occurrence

        # Step 2: Insert a New Table Just Before the "Signature" Paragraph
        new_table = doc.add_table(rows=0, cols=2)
        new_table.style = "Table Grid"

        for item in qa_data:
            question = item.get("question", "").strip()
            answer = item.get("answer", "").strip()

            row = new_table.add_row().cells
            row[0].text = question
            row[1].text = answer

        if signature_paragraph:
            # Find the XML position of "Signature" and insert the table
            signature_element = signature_paragraph._element
            signature_element.addprevious(new_table._element)

            # Step 3: Add a Line Break Before "Signature"
            line_break = doc.add_paragraph()  # Empty paragraph for spacing
            signature_element.addprevious(line_break._element)

        else:
            print("WARNING: 'Signature' section not found, adding table at the end.")
            doc.add_paragraph()  # Ensure spacing
            doc._element.body.append(new_table._element)

        # Step 4: Apply Formatting to the Table
        for row in new_table.rows:
            if len(row.cells) >= 2:
                col1, col2 = row.cells[0], row.cells[1]

                # Get cleaned text from column 2
                col2_text = col2.text.strip().lower()

                # If column 2 is exactly "section", clear it and merge
                if col2_text == "section":
                    col2.text = ""  # Clears column 2
                    col1.merge(col2)
                    col1.text = col1.text.strip()

                    # Apply formatting only to merged rows
                    for paragraph in col1.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(14)
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)  # White text

                    # Apply dark blue background
                    shading_xml = parse_xml(r'<w:shd {} w:fill="002060"/>'.format(nsdecls('w')))
                    col1._element.get_or_add_tcPr().append(shading_xml)

        # Step 5: Save modified DOCX to a buffer
        output_docx = io.BytesIO()
        doc.save(output_docx)
        output_docx.seek(0)

        return output_docx.read(), "Success"

    except Exception as e:
        return None, f"Processing error: {str(e)}"